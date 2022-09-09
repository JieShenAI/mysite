
from django.shortcuts import render, redirect, HttpResponse
from django.views.decorators.csrf import csrf_exempt
import json
from pathlib import Path
from django.http import JsonResponse
from http import HTTPStatus

# TODO 上线记得开启
from app.views.jsons import _saveImg, get_file_json_by_fid

import os
from django.http import Http404, StreamingHttpResponse, FileResponse
from docx import Document

"""
文件相关的操作：
    下载规划word文档
"""

""""
以下是Vue的文档导出
"""


@csrf_exempt
def downWord(request):
    """
        接受用户post传递上来的参数，再将其中的文字，拼成一个word文档，返回给用户
    """
    if request.method == "GET":
        return HttpResponse("You should access by POST.")

    import simplejson
    if request.method == "POST":
        data = simplejson.loads(request.body)
        _generateWord(data)
        fileLocation = "./test.docx"
        try:
            response = FileResponse(open(fileLocation, 'rb'))
            response['content_type'] = "application/octet-stream"
            response['Content-Disposition'] = 'attachment; filename=' + \
                os.path.basename(fileLocation)
            return response
        except Exception:
            raise Http404


def _generateWord(dict_arr):
    """
        处理用户POST传递上来的数据
    """
    document = Document()
    data = list(dict_arr.items())

    def sortKey(key):
        a, b = key.split('-')
        return 10*int(a)+int(b)
    data.sort(key=lambda d: sortKey(d[0]))
    for arrObj in data:
        for obj in arrObj[1]:
            key, value = list(obj.items())[0]
            if key == 'P':
                document.add_paragraph(value)
            elif key[0] == 'H':
                document.add_heading(value, level=int(key[1]))
            elif key == "IMG":
                p = Path(".").joinpath("app").joinpath(
                    value[value.index("static")::])
                picture = document.add_picture(str(p.absolute()))
                picture.width = int(
                    picture.width * 0.80)  # 宽度缩放为原来的30%
                picture.height = int(
                    picture.height * 0.80)  # 高度缩放为原来的30%
    document.save("./test.docx")


# D:\github\mysite + "\app"


if __name__ == "__main__":
    pass
